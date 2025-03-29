#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Google 文档下载工具
用于下载服务账号有权限访问的 Google 文档和表格内容
"""

import os
import logging
import sys
import json
import csv
import io
import time
import openpyxl
from openpyxl import Workbook
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from config import SCOPES, SERVICE_ACCOUNT_FILE

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# 创建下载目录
DOWNLOAD_DIR = 'downloads'
if not os.path.exists(DOWNLOAD_DIR):
    os.makedirs(DOWNLOAD_DIR)

class GoogleDocsDownloader:
    """Google 文档下载工具类"""
    
    def __init__(self, credentials_file=None, rate_limit_delay=1.5):
        """初始化下载工具"""
        self.credentials_file = credentials_file or SERVICE_ACCOUNT_FILE
        self.credentials = None
        self.drive_service = None
        self.sheets_service = None
        self.docs_service = None
        # 请求速率限制（秒）
        self.rate_limit_delay = rate_limit_delay
        
    def connect(self):
        """连接到 Google API"""
        try:
            # 检查凭证文件是否存在
            if not os.path.exists(self.credentials_file):
                logging.error(f"凭证文件 {self.credentials_file} 不存在")
                return False
                
            logging.info(f"正在加载凭证文件: {self.credentials_file}")
            
            # 创建凭证
            self.credentials = service_account.Credentials.from_service_account_file(
                self.credentials_file, scopes=SCOPES)
            
            # 输出服务账号邮箱
            logging.info(f"服务账号邮箱: {self.credentials.service_account_email}")
            
            # 创建 Drive 服务 (用于列出文件)
            self.drive_service = build('drive', 'v3', credentials=self.credentials)
            
            # 创建 Sheets 服务 (用于操作表格)
            self.sheets_service = build('sheets', 'v4', credentials=self.credentials)
            
            # 创建 Docs 服务 (用于操作文档)
            self.docs_service = build('docs', 'v1', credentials=self.credentials)
            
            logging.info("成功连接到 Google API")
            return True
            
        except Exception as e:
            logging.error(f"连接 Google API 失败: {e}")
            import traceback
            logging.error(traceback.format_exc())
            return False
    
    def list_accessible_files(self, file_type=None, max_results=100):
        """
        列出服务账号有权限访问的所有文件
        
        参数:
            file_type: 文件类型过滤，例如 'application/vnd.google-apps.spreadsheet' 表示只列出表格
            max_results: 最大结果数量
        """
        if not self.drive_service:
            if not self.connect():
                return []
                
        try:
            query = ""
            if file_type:
                query = f"mimeType='{file_type}'"
                
            logging.info(f"正在查询可访问的文件，查询条件: {query or '所有文件'}")
            
            # 列出文件
            results = self.drive_service.files().list(
                q=query,
                pageSize=max_results,
                fields="nextPageToken, files(id, name, mimeType, webViewLink)"
            ).execute()
            
            items = results.get('files', [])
            
            if not items:
                logging.info("未找到可访问的文件")
                return []
                
            logging.info(f"找到 {len(items)} 个可访问的文件")
            return items
            
        except HttpError as error:
            logging.error(f"查询文件失败: {error}")
            return []
    
    def download_spreadsheet(self, spreadsheet_id, output_format='excel'):
        """
        下载指定的表格
        
        参数:
            spreadsheet_id: 表格ID
            output_format: 输出格式，支持 'csv'、'json' 或 'excel'
        """
        if not self.sheets_service:
            if not self.connect():
                return False
                
        try:
            # 获取表格信息
            spreadsheet = self.sheets_service.spreadsheets().get(
                spreadsheetId=spreadsheet_id
            ).execute()
            
            spreadsheet_name = spreadsheet.get('properties', {}).get('title', 'unknown')
            sheets = spreadsheet.get('sheets', [])
            
            logging.info(f"正在下载表格: {spreadsheet_name} (ID: {spreadsheet_id})")
            
            # 如果是 Excel 格式，创建一个工作簿对象
            if output_format == 'excel':
                workbook = Workbook()
                # 删除默认创建的工作表
                if 'Sheet' in workbook.sheetnames:
                    del workbook['Sheet']
            
            # 为每个工作表创建一个文件
            for sheet in sheets:
                sheet_name = sheet.get('properties', {}).get('title', 'Sheet')
                sheet_id = sheet.get('properties', {}).get('sheetId', 0)
                
                # 获取工作表数据
                result = self.sheets_service.spreadsheets().values().get(
                    spreadsheetId=spreadsheet_id,
                    range=sheet_name
                ).execute()
                
                # 添加延迟，避免触发 API 配额限制
                time.sleep(self.rate_limit_delay)
                
                values = result.get('values', [])
                
                if not values:
                    logging.info(f"工作表 '{sheet_name}' 为空")
                    continue
                
                # 创建安全的文件名
                safe_spreadsheet_name = self._safe_filename(spreadsheet_name)
                safe_sheet_name = self._safe_filename(sheet_name)
                
                # 保存数据
                if output_format == 'csv':
                    filename = os.path.join(DOWNLOAD_DIR, f"{safe_spreadsheet_name}_{safe_sheet_name}.csv")
                    with open(filename, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        writer.writerows(values)
                    logging.info(f"已保存工作表 '{sheet_name}' 到 {filename}")
                    
                elif output_format == 'json':
                    filename = os.path.join(DOWNLOAD_DIR, f"{safe_spreadsheet_name}_{safe_sheet_name}.json")
                    
                    # 如果有标题行，使用它作为键
                    if len(values) > 0:
                        headers = values[0]
                        data = []
                        
                        for row in values[1:]:
                            # 确保行长度与标题行相同
                            row_data = {}
                            for i, value in enumerate(row):
                                if i < len(headers):
                                    row_data[headers[i]] = value
                                else:
                                    break
                            data.append(row_data)
                            
                        with open(filename, 'w', encoding='utf-8') as f:
                            json.dump(data, f, ensure_ascii=False, indent=2)
                        logging.info(f"已保存工作表 '{sheet_name}' 到 {filename}")
                    else:
                        logging.info(f"工作表 '{sheet_name}' 为空或没有标题行")
                        
                elif output_format == 'excel':
                    # 创建工作表（Excel工作表名称最长31个字符）
                    safe_sheet_name = sheet_name[:31]
                    # 避免重复的工作表名称
                    if safe_sheet_name in workbook.sheetnames:
                        i = 1
                        while f"{safe_sheet_name}_{i}" in workbook.sheetnames and i < 100:
                            i += 1
                        safe_sheet_name = f"{safe_sheet_name}_{i}"
                    
                    ws = workbook.create_sheet(title=safe_sheet_name)
                    
                    # 写入数据
                    for row_idx, row_data in enumerate(values):
                        for col_idx, cell_value in enumerate(row_data):
                            ws.cell(row=row_idx+1, column=col_idx+1, value=cell_value)
                    
                    logging.info(f"已添加工作表 '{sheet_name}' 到Excel文件")
            
            # 如果是Excel格式，保存整个工作簿
            if output_format == 'excel':
                excel_filename = os.path.join(DOWNLOAD_DIR, f"{safe_spreadsheet_name}.xlsx")
                workbook.save(excel_filename)
                logging.info(f"已保存Excel文件到 {excel_filename}")
            
            return True
            
        except HttpError as error:
            logging.error(f"下载表格失败: {error}")
            return False
    
    def download_document(self, document_id, output_format='docx'):
        """
        下载指定的文档
        
        参数:
            document_id: 文档ID
            output_format: 输出格式，支持 'txt' 或 'docx'
        """
        if not self.docs_service:
            if not self.connect():
                return False
                
        try:
            # 获取文档信息
            document = self.docs_service.documents().get(
                documentId=document_id
            ).execute()
            
            document_name = document.get('title', 'unknown')
            content = document.get('body', {}).get('content', [])
            
            logging.info(f"正在下载文档: {document_name} (ID: {document_id})")
            
            # 创建安全的文件名
            safe_document_name = self._safe_filename(document_name)
            
            if output_format == 'txt':
                # 提取文本内容
                text_content = self._extract_document_text(content)
                
                # 保存文本内容
                filename = os.path.join(DOWNLOAD_DIR, f"{safe_document_name}.txt")
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                    
                logging.info(f"已保存文档 '{document_name}' 到 {filename}")
                return True
            
            elif output_format == 'docx':
                # 创建Word文档
                doc = Document()
                doc.add_heading(document_name, 0)
                
                # 处理文档内容
                for item in content:
                    if 'paragraph' in item:
                        paragraph = item['paragraph']
                        elements = paragraph.get('elements', [])
                        
                        paragraph_text = []
                        for element in elements:
                            if 'textRun' in element:
                                paragraph_text.append(element['textRun'].get('content', ''))
                        
                        if ''.join(paragraph_text).strip():
                            doc.add_paragraph(''.join(paragraph_text))
                    
                    # 处理表格 (简化处理)
                    elif 'table' in item:
                        table = doc.add_table(rows=1, cols=1)
                        cell = table.cell(0, 0)
                        cell.text = "[表格内容]" # 简化处理，实际应用中可以更复杂地处理表格
                    
                    # 处理分节符
                    elif 'sectionBreak' in item:
                        doc.add_page_break()
                
                # 保存Word文档
                filename = os.path.join(DOWNLOAD_DIR, f"{safe_document_name}.docx")
                doc.save(filename)
                
                logging.info(f"已保存文档 '{document_name}' 到 {filename}")
                return True
            
        except HttpError as error:
            logging.error(f"下载文档失败: {error}")
            return False
    
    def _extract_document_text(self, content):
        """从文档内容中提取文本"""
        text = []
        
        for item in content:
            if 'paragraph' in item:
                paragraph = item['paragraph']
                elements = paragraph.get('elements', [])
                
                paragraph_text = []
                for element in elements:
                    if 'textRun' in element:
                        paragraph_text.append(element['textRun'].get('content', ''))
                
                text.append(''.join(paragraph_text))
            
            # 处理表格
            elif 'table' in item:
                table = item['table']
                table_rows = table.get('tableRows', [])
                
                for row in table_rows:
                    row_text = []
                    cells = row.get('tableCells', [])
                    
                    for cell in cells:
                        cell_content = cell.get('content', [])
                        cell_text = self._extract_document_text(cell_content)
                        row_text.append(cell_text)
                    
                    text.append(' | '.join(row_text))
            
            # 处理列表
            elif 'sectionBreak' in item:
                text.append('\n---\n')
        
        return '\n'.join(text)
    
    def _safe_filename(self, filename):
        """创建安全的文件名"""
        # 替换不安全的字符
        return "".join([c if c.isalnum() or c in ['-', '_', '.', ' '] else '_' for c in filename])


def main():
    """主函数"""
    # 解析命令行参数
    import argparse
    parser = argparse.ArgumentParser(description='Google 文档下载工具')
    parser.add_argument('--list', action='store_true', help='列出可访问的文件')
    parser.add_argument('--download-sheet', metavar='SHEET_ID', help='下载指定的表格')
    parser.add_argument('--download-doc', metavar='DOC_ID', help='下载指定的文档')
    parser.add_argument('--format', choices=['csv', 'json', 'excel', 'docx', 'txt'], default='excel', 
                      help='下载格式 (默认: excel，表格支持excel/csv/json，文档支持docx/txt)')
    parser.add_argument('--download-all', action='store_true', help='下载所有可访问的文件')
    parser.add_argument('--delay', type=float, default=1.5, help='API请求间隔延迟，单位秒 (默认: 1.5)')
    
    args = parser.parse_args()
    
    # 创建下载器实例，设置请求延迟
    downloader = GoogleDocsDownloader(rate_limit_delay=args.delay)
    
    if not downloader.connect():
        logging.error("连接 Google API 失败")
        return 1
    

    
    # 列出可访问的文件
    if args.list:
        # 列出可访问的表格
        print("\n=== 可访问的表格 ===")
        spreadsheets = downloader.list_accessible_files(
            file_type='application/vnd.google-apps.spreadsheet'
        )
        
        if spreadsheets:
            print(f"\n找到 {len(spreadsheets)} 个可访问的表格:")
            print("-" * 80)
            print(f"{'序号':<5} {'表格名':<40} {'表格ID':<40}")
            print("-" * 80)
            
            for i, sheet in enumerate(spreadsheets, 1):
                print(f"{i:<5} {sheet.get('name', '')[:38]:<40} {sheet.get('id', ''):<40}")
        else:
            print("未找到可访问的表格")
        
        # 列出可访问的文档
        print("\n=== 可访问的文档 ===")
        documents = downloader.list_accessible_files(
            file_type='application/vnd.google-apps.document'
        )
        
        if documents:
            print(f"\n找到 {len(documents)} 个可访问的文档:")
            print("-" * 80)
            print(f"{'序号':<5} {'文档名':<40} {'文档ID':<40}")
            print("-" * 80)
            
            for i, doc in enumerate(documents, 1):
                print(f"{i:<5} {doc.get('name', '')[:38]:<40} {doc.get('id', ''):<40}")
        else:
            print("未找到可访问的文档")
    
    # 下载指定的表格
    if args.download_sheet:
        if args.format in ['csv', 'json', 'excel']:
            success = downloader.download_spreadsheet(args.download_sheet, args.format)
            if success:
                print(f"表格下载完成，已保存到 {DOWNLOAD_DIR} 目录")
            else:
                print("表格下载失败")
        else:
            print(f"错误：表格下载不支持 {args.format} 格式，请使用 csv、json 或 excel 格式")
            return 1
    
    # 下载指定的文档
    if args.download_doc:
        if args.format in ['txt', 'docx']:
            success = downloader.download_document(args.download_doc, args.format)
            if success:
                print(f"文档下载完成，已保存到 {DOWNLOAD_DIR} 目录")
            else:
                print("文档下载失败")
        else:
            print(f"错误：文档下载不支持 {args.format} 格式，请使用 txt 或 docx 格式")
            return 1
    
    # 下载所有可访问的文件
    if args.download_all:
        # 下载表格
        spreadsheets = downloader.list_accessible_files(
            file_type='application/vnd.google-apps.spreadsheet'
        )
        
        if spreadsheets:
            print(f"\n开始下载 {len(spreadsheets)} 个表格...")
            for i, sheet in enumerate(spreadsheets, 1):
                sheet_id = sheet.get('id')
                sheet_name = sheet.get('name')
                print(f"[{i}/{len(spreadsheets)}] 正在下载表格: {sheet_name}")
                if args.format in ['csv', 'json', 'excel']:
                    downloader.download_spreadsheet(sheet_id, args.format)
                else:
                    # 默认使用Excel格式
                    downloader.download_spreadsheet(sheet_id, 'excel')
        
        # 下载文档
        documents = downloader.list_accessible_files(
            file_type='application/vnd.google-apps.document'
        )
        
        if documents:
            print(f"\n开始下载 {len(documents)} 个文档...")
            for i, doc in enumerate(documents, 1):
                doc_id = doc.get('id')
                doc_name = doc.get('name')
                print(f"[{i}/{len(documents)}] 正在下载文档: {doc_name}")
                if args.format in ['txt', 'docx']:
                    downloader.download_document(doc_id, args.format)
                else:
                    # 默认使用docx格式
                    downloader.download_document(doc_id, 'docx')
        
        print(f"\n所有文件下载完成，已保存到 {DOWNLOAD_DIR} 目录")
    
    # 如果没有指定任何操作，显示帮助信息
    if not (args.list or args.download_sheet or args.download_doc or args.download_all):
        parser.print_help()
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
